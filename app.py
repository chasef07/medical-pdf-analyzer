import gradio as gr
import base64
import anthropic
from anthropic.types.message_create_params import MessageCreateParamsNonStreaming
from anthropic.types.messages.batch_create_params import Request
from config import ANTHROPIC_API_KEY
import os
import tempfile
import time
import json
from typing import List, Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_structured_data_from_pdf(pdf_file, client, file_index):
    """
    Step 1: Extract structured JSON data from a single PDF
    """
    # Check file size before processing (256MB = 268,435,456 bytes)
    file_size = os.path.getsize(pdf_file)
    max_size = 200 * 1024 * 1024  # 200MB to be safe with base64 encoding overhead

    if file_size > max_size:
        return {
            "error": f"Document {file_index + 1} ({os.path.basename(pdf_file)}) is too large: {file_size / (1024*1024):.1f}MB. Maximum size is {max_size / (1024*1024):.0f}MB.",
            "file_size": file_size
        }

    with open(pdf_file, "rb") as f:
        pdf_data = base64.standard_b64encode(f.read()).decode("utf-8")

    json_extraction_prompt = f"""You are a medical-legal data extraction specialist. Extract structured data from this medical document ({os.path.basename(pdf_file)}) into valid JSON format.

CRITICAL: Output ONLY valid JSON. Quote exact text for all medical findings.

Required JSON structure:
{{
  "document_metadata": {{
    "filename": "{os.path.basename(pdf_file)}",
    "analysis_date": "{datetime.now().isoformat()}",
    "document_type": "initial_visit|follow_up|imaging|billing|legal|other"
  }},
  "patient_info": {{
    "demographics": {{
      "name": "string or null",
      "age": "string or null",
      "gender": "string or null",
      "height": "string or null",
      "weight": "string or null"
    }},
    "incident_details": {{
      "date": "YYYY-MM-DD or null",
      "time": "string or null",
      "mechanism": "string or null",
      "location": "string or null",
      "circumstances": "string or null",
      "source_quote": "exact text from document"
    }}
  }},
  "medical_encounters": [
    {{
      "date": "YYYY-MM-DD",
      "provider": "string",
      "facility": "string",
      "visit_type": "string",
      "chief_complaints": ["list of complaints"],
      "examination_findings": {{
        "orthopedic_tests": {{
          "compression_test": "positive/negative",
          "jacksons_test": "positive/negative",
          "distraction_test": "positive/negative",
          "shoulder_depression_test": "positive/negative",
          "kemps_test": "positive/negative",
          "spurlings_test": "positive/negative/not_performed",
          "hoffmans_test": "positive/negative/not_performed",
          "other_tests": ["test_name: result"]
        }},
        "range_of_motion": {{
          "cervical": "normal/decreased/restricted_with_pain",
          "thoracic": "normal/decreased/restricted_with_pain",
          "lumbar": "normal/decreased/restricted_with_pain",
          "percentage_decrease": "percentage if specified",
          "details": ["specific ROM measurements"]
        }},
        "neurological": {{
          "motor_strength": {{"C5": "0-5/5", "C6": "0-5/5", "C7": "0-5/5", "C8": "0-5/5", "T1": "0-5/5"}},
          "reflexes": ["reflex: result"],
          "sensation": ["area: intact/diminished/absent"],
          "cranial_nerves": "intact/abnormal"
        }},
        "palpation": {{
          "tenderness": ["cervical/thoracic/lumbar: present/absent"],
          "spasm": ["cervical/thoracic/lumbar: present/absent"],
          "trigger_points": ["location"]
        }},
        "vital_signs": {{
          "blood_pressure": "systolic/diastolic",
          "pulse": "bpm",
          "temperature": "degrees F",
          "respiratory_rate": "breaths/min",
          "oxygen_saturation": "percentage",
          "height": "measurement",
          "weight": "measurement"
        }}
      }},
      "diagnostic_results": {{
        "imaging": {{
          "study_type": "X-ray/MRI/CT",
          "study_date": "YYYY-MM-DD",
          "ordering_provider": "provider name",
          "radiologist": "radiologist name",
          "findings_detailed": ["bullet point finding with exact quote"],
          "impression_detailed": ["bullet point impression with exact quote"],
          "technical_details": ["axial image numbers, measurements, etc."]
        }},
        "nerve_conduction_studies": {{
          "study_date": "YYYY-MM-DD",
          "provider": "provider name",
          "motor_findings": ["specific findings"],
          "sensory_findings": ["specific findings"],
          "f_wave_findings": ["specific findings"],
          "conclusion": "exact conclusion text"
        }},
        "lab_tests": ["test: result"],
        "other_studies": ["study: findings"]
      }},
      "treatments_provided": ["treatment list"],
      "provider_opinions": [
        {{
          "opinion": "exact quote",
          "source_quote": "verbatim text",
          "opinion_type": "causation|prognosis|treatment"
        }}
      ]
    }}
  ],
  "billing_data": [
    {{
      "date": "YYYY-MM-DD",
      "provider": "string",
      "facility": "string",
      "service_code": "string",
      "service_description": "string",
      "amount_charged": "number or null",
      "balance_due": "number or null",
      "insurance_payments": "number or null",
      "billing_observations": "any unusual charges or observations",
      "source_quote": "exact text"
    }}
  ],
  "quality_flags": {{
    "inconsistencies": ["description of issue"],
    "missing_data": ["what data is missing"],
    "date_discrepancies": ["date conflict description"],
    "unclear_text": ["illegible or uncertain sections"]
  }}
}}

VERIFICATION RULES:
- Include source quotes for every medical finding
- Use null for missing data, not empty strings
- Flag any illegible text as "unclear_text"
- Mark any assumptions as "inferred_data" in quality_flags
- Extract exact measurements and dates
- Preserve medical terminology as written

Output ONLY the JSON, no other text."""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_data
                            }
                        },
                        {
                            "type": "text",
                            "text": json_extraction_prompt
                        }
                    ]
                }
            ]
        )

        json_text = response.content[0].text

        # Enhanced cleaning of potential markdown formatting and common issues
        json_text = json_text.strip()

        # Remove markdown code blocks
        if json_text.startswith('```json'):
            json_text = json_text.replace('```json', '').replace('```', '').strip()
        elif json_text.startswith('```'):
            json_text = json_text.replace('```', '').strip()

        # Remove any leading/trailing text that isn't JSON
        start_brace = json_text.find('{')
        end_brace = json_text.rfind('}')
        if start_brace != -1 and end_brace != -1 and end_brace > start_brace:
            json_text = json_text[start_brace:end_brace + 1]

        try:
            return json.loads(json_text)
        except json.JSONDecodeError:
            # If JSON parsing still fails, try to handle as text extraction
            return {
                "error": f"JSON parsing failed for document {file_index + 1}, attempting text fallback",
                "raw_response": response.content[0].text[:1000] + "..." if len(response.content[0].text) > 1000 else response.content[0].text,
                "extracted_text": json_text[:500] + "..." if len(json_text) > 500 else json_text,
                "fallback_summary": f"Document {file_index + 1} processed but returned non-JSON format"
            }

    except json.JSONDecodeError as e:
        return {
            "error": f"Failed to parse JSON from document {file_index + 1}",
            "raw_response": response.content[0].text[:1000] + "..." if len(response.content[0].text) > 1000 else response.content[0].text,
            "json_error": str(e)
        }
    except Exception as e:
        return {
            "error": f"Failed to extract data from document {file_index + 1}",
            "exception": str(e),
            "filename": os.path.basename(pdf_file)
        }

def synthesize_unified_analysis(json_extractions, client):
    """
    Step 2: Create a doctor-style chronological medical summary
    """
    # Filter out error documents but keep track of them
    successful_extractions = []
    failed_extractions = []

    for i, extraction in enumerate(json_extractions):
        if "error" in extraction:
            failed_extractions.append({
                "index": i + 1,
                "error": extraction.get("error", "Unknown error"),
                "filename": extraction.get("filename", f"Document {i + 1}")
            })
        else:
            successful_extractions.append(extraction)

    if not successful_extractions:
        return "Unable to process any documents successfully. All documents had extraction errors."

    synthesis_prompt = f"""You are a medical professional writing a chronological record review. Create a professional medical record summary that EXACTLY matches the style and format of the gold standard medical record review shown below.

EXTRACTED DATA FROM {len(successful_extractions)} SUCCESSFULLY PROCESSED DOCUMENTS:
{json.dumps(successful_extractions, indent=2)}

GOLD STANDARD FORMAT EXAMPLE:
Record Review:
10/05/2024, Florida Crash Report – John Doe was a restrained driver when another vehicle reversed into his driveway and impacted John Doe's vehicle. Airbags did not deploy, and rescue was declined.

10/06/2024, Silverman Chiropractic & Rehabilitation Center, Gregory Mazzotta, DC, Initial Evaluation, report states John Doe was hit on the left side. His initial complaints were acute cervical pain, acute thoracic pain, and headache. Physical exam consisted of positive compression, positive Jackson's, positive distraction, positive shoulder depression, and positive Kemp's (thoracic), as well as palpable spasm with restricted motion in cervical and thoracic spines.

Of note, report incorrectly states Date of Accident as 10/07/2024. Additionally, no mention of +/- radiculopathy or any strength testing during initial exam.

10/11/2024, Central Magnetic Imaging, Heather Kahan, MD, Neuroradiologist
MRI Cervical Spine revealed:

FINDINGS:
• The vertebral bodies are in normal anatomic alignment. Normal vertebral stature is maintained. There is marked straightening of the normal cervical lordosis.
• At the level of C3-C4 there is a 1mm broad-based right paracentral disc herniation demonstrated on axial image #24 and 25 with compression on the thecal sac and partial disc desiccation.

IMPRESSION:
• There is marked straightening of the normal cervical lordosis consistent with marked torticollis/cervical strain.

CRITICAL FORMATTING RULES:
1. **START WITH**: "Record Review:" (no other text before this)
2. **INDIVIDUAL ORTHO TESTS**: List each test separately ("positive compression, positive Jackson's, positive distraction")
3. **IMAGING STUDIES**: Use bullet points for FINDINGS and IMPRESSION sections
4. **MULTIPLE VISIT DATES**: Group same-provider routine visits on one line: "10/08/2024, 10/10/2024, 10/14/2024, Provider, treatment consisting of..."
5. **"Of note" THROUGHOUT**: Add "Of note," observations immediately after relevant entries (not just at end)
6. **MEDICAL ABBREVIATIONS**: Use "DC", "MD", "CMT", "EMS", "ROM" etc.
7. **EXACT QUOTES**: Include technical details like "axial image #24" and specific measurements
8. **PROVIDER CREDENTIALS**: Always include credentials (MD, DC, APRN, etc.)

SPECIFIC INSTRUCTIONS:
- Extract individual orthopedic test results (compression, Jackson's, distraction, Spurling's, etc.)
- For imaging studies, create full FINDINGS and IMPRESSION sections with bullet points
- Group multiple routine treatment visits on single lines
- Include "Of note," statements immediately after relevant entries
- Use medical abbreviations consistently
- Include specific technical details from imaging reports
- Add billing section with specific amounts and professional observations

OUTPUT REQUIREMENTS:
- Start immediately with "Record Review:"
- Follow exact chronological order
- Use professional medical language
- Include all technical details and measurements
- End with Medical Bills section if billing data exists"""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            messages=[
                {
                    "role": "user",
                    "content": synthesis_prompt
                }
            ]
        )

        doctor_summary = response.content[0].text

        # Add information about failed extractions if any
        if failed_extractions:
            doctor_summary += f"\n\nPROCESSING NOTES:\n"
            for failed in failed_extractions:
                doctor_summary += f"Document {failed['index']} ({failed.get('filename', 'Unknown')}) could not be processed: {failed['error']}\n"

        return doctor_summary

    except Exception as e:
        return f"Error creating medical summary: {str(e)}"

def cross_reference_findings(json_extractions):
    """
    Utility function to perform cross-referencing logic
    """
    correlations = {
        "timeline_conflicts": [],
        "diagnostic_progression": [],
        "treatment_evolution": [],
        "billing_anomalies": []
    }

    # Extract all dates and check for conflicts
    all_encounters = []
    for doc_idx, extraction in enumerate(json_extractions):
        if "medical_encounters" in extraction:
            for encounter in extraction["medical_encounters"]:
                encounter["source_document"] = doc_idx + 1
                all_encounters.append(encounter)

    # Sort chronologically
    all_encounters.sort(key=lambda x: x.get("date", "9999-12-31"))

    # Check for same-date conflicts
    date_groups = {}
    for encounter in all_encounters:
        date = encounter.get("date")
        if date:
            if date not in date_groups:
                date_groups[date] = []
            date_groups[date].append(encounter)

    for date, encounters in date_groups.items():
        if len(encounters) > 1:
            correlations["timeline_conflicts"].append({
                "date": date,
                "encounters": encounters,
                "issue": "Multiple encounters on same date"
            })

    return correlations

def generate_docx_report(medical_summary, pdf_files, successful_count):
    """
    Generate a professional Word document from the medical summary
    """
    # Create a new Document
    doc = Document()

    # Add title
    title = doc.add_heading('Medical Record Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add header information
    header_para = doc.add_paragraph()
    header_para.add_run('Generated: ').bold = True
    header_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    header_para.add_run('\nDocuments Analyzed: ').bold = True
    header_para.add_run(f'{successful_count} of {len(pdf_files)} processed successfully')

    # Add a line break
    doc.add_paragraph()

    # Add the medical summary content
    # Split the summary into sections for better formatting
    lines = medical_summary.split('\n')
    current_paragraph = None

    for line in lines:
        line = line.strip()
        if not line:
            if current_paragraph:
                doc.add_paragraph()
            current_paragraph = None
            continue

        # Check if it's a section header
        if line.startswith('Record Review:'):
            heading = doc.add_heading(line, level=1)
        elif line.startswith('Medical Bills:'):
            heading = doc.add_heading(line, level=1)
        elif line.startswith('FINDINGS:') or line.startswith('IMPRESSION:'):
            heading = doc.add_heading(line, level=2)
        elif line.startswith('Of note,'):
            para = doc.add_paragraph()
            para.add_run(line).italic = True
        elif line.startswith('•'):
            # Bullet point
            para = doc.add_paragraph(line, style='List Bullet')
        else:
            # Regular paragraph
            if not current_paragraph:
                current_paragraph = doc.add_paragraph()
            if current_paragraph.text:
                current_paragraph.add_run('\n')
            current_paragraph.add_run(line)

    # Add document processing summary
    doc.add_heading('Document Processing Summary', level=1)
    for i, pdf_file in enumerate(pdf_files):
        status = "Processed" if i < successful_count else "Failed"
        para = doc.add_paragraph()
        para.add_run(f'Document {i+1}: ').bold = True
        para.add_run(f'{os.path.basename(pdf_file)} - {status}')

    # Save to timestamped file
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"medical_analysis_report_{timestamp}.docx"
    doc.save(filename)

    return filename

def analyze_medical_pdf(pdf_files):
    """
    Two-step analysis: Extract structured JSON data, then synthesize unified timeline
    """
    if pdf_files is None or len(pdf_files) == 0:
        return "Please upload at least one PDF file.", None

    try:
        # Initialize the Anthropic client
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        # Step 1: Extract structured JSON data from each PDF
        json_extractions = []
        extraction_status = f"TWO-STEP ANALYSIS PROCESS\n\nStep 1: Extracting structured data from {len(pdf_files)} documents...\n\n"

        for i, pdf_file in enumerate(pdf_files):
            extraction_status += f"Processing {os.path.basename(pdf_file)}... "

            json_data = extract_structured_data_from_pdf(pdf_file, client, i)
            json_extractions.append(json_data)

            if "error" in json_data:
                extraction_status += f"Error\n"
            else:
                extraction_status += f"Complete\n"

        # Check if all extractions failed
        successful_extractions = [data for data in json_extractions if "error" not in data]
        failed_extractions = [data for data in json_extractions if "error" in data]

        if not successful_extractions:
            error_report = "\n\nALL EXTRACTION FAILED:\n"
            for i, error_data in enumerate(failed_extractions):
                error_report += f"- Document {i+1}: {error_data.get('error', 'Unknown error')}\n"
            return extraction_status + error_report, None

        # Step 2: Generate doctor-style medical summary
        extraction_status += "\nStep 2: Creating professional medical record summary...\nComplete!\n\n"

        # Generate doctor-style summary
        medical_summary = synthesize_unified_analysis(json_extractions, client)

        # Create complete markdown report
        markdown_report = f"""# MEDICAL RECORD ANALYSIS
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Documents Analyzed: {len(successful_extractions)} of {len(pdf_files)} processed successfully

{extraction_status}{'='*80}

{medical_summary}

{'='*80}

## DOCUMENT PROCESSING SUMMARY:
"""
        for i, pdf_file in enumerate(pdf_files):
            status = "Processed" if i < len(json_extractions) and "error" not in json_extractions[i] else "Failed"
            markdown_report += f"Document {i+1}: {os.path.basename(pdf_file)} - {status}\n"

        # Save markdown backup
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        markdown_filename = f"medical_analysis_{timestamp}.md"
        with open(markdown_filename, 'w', encoding='utf-8') as f:
            f.write(markdown_report)

        # Generate DOCX file
        docx_file_path = generate_docx_report(medical_summary, pdf_files, len(successful_extractions))

        # Create status message
        status_message = f"""Analysis Complete!

{extraction_status}

Medical report generated successfully.
Documents processed: {len(successful_extractions)} of {len(pdf_files)}

SAVED FILES:
- Markdown backup: {markdown_filename}
- Word document: Ready for download

The professional medical record analysis has been generated and is ready for download."""

        return markdown_report, docx_file_path

    except Exception as e:
        # Enhanced error handling for batch API
        error_msg = f"Error analyzing PDFs with batch API: {str(e)}\n\n"

        if "batch" in str(e).lower():
            error_msg += "This appears to be a batch processing error. Please check:\n"
            error_msg += "- Your API key has batch processing permissions\n"
            error_msg += "- The total file size doesn't exceed 256MB\n"
            error_msg += "- All uploaded files are valid PDFs\n"
        elif "413" in str(e) or "request_too_large" in str(e).lower():
            error_msg += "Request too large error. The combined PDF files exceed the 256MB limit.\n"
            error_msg += "Please try uploading fewer or smaller PDF files.\n"
        elif "rate" in str(e).lower() or "limit" in str(e).lower():
            error_msg += "Rate limit exceeded. Please wait a moment and try again.\n"
        elif "api_key" in str(e).lower() or "authentication" in str(e).lower():
            error_msg += "API authentication error. Please check your Anthropic API key.\n"

        return error_msg, None

# Create the Gradio interface
with gr.Blocks(title="Medical PDF Analyzer") as demo:
    gr.Markdown("# Medical PDF Legal Analysis Tool")
    gr.Markdown("Upload multiple medical PDF documents to generate a unified timeline analysis with cross-document correlation. Uses two-step processing: structured extraction and synthesis.")
    
    with gr.Row():
        with gr.Column():
            pdf_input = gr.File(
                label="Upload Medical PDFs",
                file_types=[".pdf"],
                type="filepath",
                file_count="multiple"
            )
            analyze_btn = gr.Button("Analyze Documents")
        
        with gr.Column():
            analysis_output = gr.Textbox(
                label="Analysis Results",
                lines=30,
                max_lines=50,
                show_copy_button=True,
                placeholder="Your medical analysis will appear here..."
            )
            download_docx = gr.File(
                label="Download DOCX Report",
                file_count="single",
                file_types=[".docx"]
            )
    
    # Set up the click event
    analyze_btn.click(
        fn=analyze_medical_pdf,
        inputs=[pdf_input],
        outputs=[analysis_output, download_docx]
    )

if __name__ == "__main__":
    demo.launch(share=True)